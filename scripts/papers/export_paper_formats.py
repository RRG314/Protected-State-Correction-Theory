#!/usr/bin/env python3
"""Export final paper markdown files to TEX, DOCX, and PDF.

This exporter is intentionally lightweight and deterministic so it can run in
environments without Pandoc/LaTeX toolchains.
"""

from __future__ import annotations

import re
from pathlib import Path
from xml.sax.saxutils import escape as xml_escape

from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.utils import ImageReader
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, Preformatted, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "papers" / "formats"
PAPERS = [
    ROOT / "papers" / "recoverability_paper_final.md",
    ROOT / "papers" / "ocp_core_paper.md",
    ROOT / "papers" / "bridge_paper.md",
    ROOT / "papers" / "mhd_paper_upgraded.md",
]


IMAGE_RE = re.compile(r"^!\[(.*?)\]\((.*?)\)\s*$")
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
CODE_INLINE_RE = re.compile(r"`([^`]+)`")


def clean_inline(text: str) -> str:
    text = CODE_INLINE_RE.sub(r"\1", text)
    text = LINK_RE.sub(r"\1 (\2)", text)
    text = text.replace("**", "").replace("*", "")
    return text.strip()


def escape_latex(text: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    out = []
    for ch in text:
        out.append(replacements.get(ch, ch))
    return "".join(out)


def markdown_to_latex(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    tex: list[str] = [
        r"\documentclass[11pt]{article}",
        r"\usepackage[margin=1in]{geometry}",
        r"\usepackage[T1]{fontenc}",
        r"\usepackage{lmodern}",
        r"\usepackage{hyperref}",
        r"\usepackage{graphicx}",
        r"\usepackage{enumitem}",
        r"\usepackage{verbatim}",
        r"\setlist{nosep}",
        r"\begin{document}",
    ]

    in_code = False
    in_itemize = False
    in_enum = False
    in_verbatim_block = False

    def close_lists() -> None:
        nonlocal in_itemize, in_enum
        if in_itemize:
            tex.append(r"\end{itemize}")
            in_itemize = False
        if in_enum:
            tex.append(r"\end{enumerate}")
            in_enum = False

    for raw in lines:
        line = raw.rstrip("\n")

        if line.strip().startswith("```"):
            close_lists()
            if not in_code:
                tex.append(r"\begin{verbatim}")
                in_code = True
            else:
                tex.append(r"\end{verbatim}")
                in_code = False
            continue

        if in_code:
            tex.append(line)
            continue

        image_match = IMAGE_RE.match(line.strip())
        if image_match:
            close_lists()
            caption = escape_latex(clean_inline(image_match.group(1)))
            image_rel = image_match.group(2)
            image_abs = (md_path.parent / image_rel).resolve()
            tex.extend(
                [
                    r"\begin{figure}[ht]",
                    r"\centering",
                    rf"\includegraphics[width=0.9\linewidth]{{\detokenize{{{image_abs.as_posix()}}}}}",
                    rf"\caption{{{caption}}}",
                    r"\end{figure}",
                ]
            )
            continue

        if line.strip().startswith("#"):
            close_lists()
            hashes, title = line.split(" ", 1)
            level = len(hashes)
            title_tex = escape_latex(clean_inline(title))
            if level == 1:
                tex.append(rf"\section*{{{title_tex}}}")
            elif level == 2:
                tex.append(rf"\subsection*{{{title_tex}}}")
            elif level == 3:
                tex.append(rf"\subsubsection*{{{title_tex}}}")
            else:
                tex.append(rf"\paragraph{{{title_tex}}}")
            continue

        if re.match(r"^\d+\.\s+", line.strip()):
            if in_itemize:
                tex.append(r"\end{itemize}")
                in_itemize = False
            if not in_enum:
                tex.append(r"\begin{enumerate}")
                in_enum = True
            item_text = re.sub(r"^\d+\.\s+", "", line.strip())
            tex.append(rf"\item {escape_latex(clean_inline(item_text))}")
            continue

        if line.strip().startswith("- "):
            if in_enum:
                tex.append(r"\end{enumerate}")
                in_enum = False
            if not in_itemize:
                tex.append(r"\begin{itemize}")
                in_itemize = True
            tex.append(rf"\item {escape_latex(clean_inline(line.strip()[2:]))}")
            continue

        if line.strip() == "":
            close_lists()
            if in_verbatim_block:
                tex.append(r"\end{verbatim}")
                in_verbatim_block = False
            tex.append("")
            continue

        if "|" in line and line.count("|") >= 2:
            close_lists()
            if not in_verbatim_block:
                tex.append(r"\begin{verbatim}")
                in_verbatim_block = True
            tex.append(line)
            continue

        if in_verbatim_block:
            tex.append(r"\end{verbatim}")
            in_verbatim_block = False

        close_lists()
        tex.append(escape_latex(clean_inline(line)))
        tex.append("")

    close_lists()
    if in_verbatim_block:
        tex.append(r"\end{verbatim}")
    if in_code:
        tex.append(r"\end{verbatim}")
    tex.append(r"\end{document}")

    out_path.write_text("\n".join(tex), encoding="utf-8")


def markdown_to_docx(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(11)

    in_code = False
    code_lines: list[str] = []
    paragraph_buf: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buf
        if paragraph_buf:
            doc.add_paragraph(clean_inline(" ".join(paragraph_buf)))
            paragraph_buf = []

    def flush_code() -> None:
        nonlocal code_lines
        if code_lines:
            p = doc.add_paragraph("\n".join(code_lines))
            for run in p.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            code_lines = []

    for raw in lines:
        line = raw.rstrip("\n")

        if line.strip().startswith("```"):
            flush_paragraph()
            if not in_code:
                in_code = True
            else:
                in_code = False
                flush_code()
            continue

        if in_code:
            code_lines.append(line)
            continue

        image_match = IMAGE_RE.match(line.strip())
        if image_match:
            flush_paragraph()
            image_path = (md_path.parent / image_match.group(2)).resolve()
            if image_path.exists():
                doc.add_picture(str(image_path), width=Inches(6.0))
                doc.add_paragraph(clean_inline(image_match.group(1)))
            continue

        if line.strip().startswith("#"):
            flush_paragraph()
            hashes, title = line.split(" ", 1)
            level = len(hashes)
            if level == 1:
                doc.add_heading(clean_inline(title), level=0)
            else:
                doc.add_heading(clean_inline(title), level=min(level - 1, 4))
            continue

        if re.match(r"^\d+\.\s+", line.strip()):
            flush_paragraph()
            item = re.sub(r"^\d+\.\s+", "", line.strip())
            doc.add_paragraph(clean_inline(item), style="List Number")
            continue

        if line.strip().startswith("- "):
            flush_paragraph()
            doc.add_paragraph(clean_inline(line.strip()[2:]), style="List Bullet")
            continue

        if line.strip() == "":
            flush_paragraph()
            continue

        paragraph_buf.append(line.strip())

    flush_paragraph()
    flush_code()
    doc.save(out_path)


def markdown_to_pdf(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = SimpleDocTemplate(str(out_path), pagesize=letter, leftMargin=54, rightMargin=54, topMargin=54, bottomMargin=54)
    styles = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=styles["Normal"], fontName="Times-Roman", fontSize=11, leading=14)
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=18, leading=22, spaceAfter=8)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=14, leading=18, spaceAfter=6)
    h3 = ParagraphStyle("H3", parent=styles["Heading3"], fontSize=12, leading=15, spaceAfter=4)
    code_style = ParagraphStyle("Code", parent=styles["Code"], fontName="Courier", fontSize=9, leading=11)
    story = []

    in_code = False
    code_lines: list[str] = []
    paragraph_buf: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buf
        if paragraph_buf:
            text = clean_inline(" ".join(paragraph_buf))
            story.append(Paragraph(xml_escape(text), body))
            story.append(Spacer(1, 6))
            paragraph_buf = []

    def flush_code() -> None:
        nonlocal code_lines
        if code_lines:
            story.append(Preformatted("\n".join(code_lines), code_style))
            story.append(Spacer(1, 6))
            code_lines = []

    for raw in lines:
        line = raw.rstrip("\n")

        if line.strip().startswith("```"):
            flush_paragraph()
            if not in_code:
                in_code = True
            else:
                in_code = False
                flush_code()
            continue

        if in_code:
            code_lines.append(line)
            continue

        image_match = IMAGE_RE.match(line.strip())
        if image_match:
            flush_paragraph()
            image_path = (md_path.parent / image_match.group(2)).resolve()
            if image_path.exists():
                img = RLImage(str(image_path))
                iw, ih = ImageReader(str(image_path)).getSize()
                max_width = 420.0
                if iw > max_width:
                    scale = max_width / float(iw)
                    img.drawWidth = iw * scale
                    img.drawHeight = ih * scale
                img.hAlign = "CENTER"
                story.append(img)
                story.append(Spacer(1, 4))
                story.append(Paragraph(xml_escape(clean_inline(image_match.group(1))), body))
                story.append(Spacer(1, 8))
            continue

        if line.strip().startswith("#"):
            flush_paragraph()
            hashes, title = line.split(" ", 1)
            level = len(hashes)
            txt = xml_escape(clean_inline(title))
            if level == 1:
                story.append(Paragraph(txt, h1))
            elif level == 2:
                story.append(Paragraph(txt, h2))
            else:
                story.append(Paragraph(txt, h3))
            story.append(Spacer(1, 4))
            continue

        if re.match(r"^\d+\.\s+", line.strip()):
            flush_paragraph()
            item = re.sub(r"^\d+\.\s+", "", line.strip())
            story.append(Paragraph(xml_escape(f"• {clean_inline(item)}"), body))
            continue

        if line.strip().startswith("- "):
            flush_paragraph()
            story.append(Paragraph(xml_escape(f"• {clean_inline(line.strip()[2:])}"), body))
            continue

        if line.strip() == "":
            flush_paragraph()
            continue

        paragraph_buf.append(line.strip())

    flush_paragraph()
    flush_code()
    doc.build(story)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for md_path in PAPERS:
        stem = md_path.stem
        tex_out = OUT_DIR / f"{stem}.tex"
        docx_out = OUT_DIR / f"{stem}.docx"
        pdf_out = OUT_DIR / f"{stem}.pdf"

        markdown_to_latex(md_path, tex_out)
        markdown_to_docx(md_path, docx_out)
        markdown_to_pdf(md_path, pdf_out)

        print(f"Exported {stem}:")
        print(f"  - {tex_out}")
        print(f"  - {docx_out}")
        print(f"  - {pdf_out}")


if __name__ == "__main__":
    main()
